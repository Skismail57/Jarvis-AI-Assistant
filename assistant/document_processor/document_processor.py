"""
Document Processor
Handles processing of PDF, Word, and Excel documents for information extraction.
"""

import os
import json
import hashlib
from typing import Optional, Dict, Any, List, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime
from enum import Enum


class DocumentType(Enum):
    PDF = "pdf"
    WORD = "word"
    EXCEL = "excel"
    TEXT = "text"
    IMAGE = "image"


class ExtractionMode(Enum):
    TEXT = "text"
    TABLES = "tables"
    METADATA = "metadata"
    FULL = "full"


@dataclass
class DocumentMetadata:
    document_id: str
    filename: str
    file_type: DocumentType
    file_size: int
    created_at: str
    processed_at: str
    page_count: int = 0
    author: Optional[str] = None
    title: Optional[str] = None
    subject: Optional[str] = None
    keywords: List[str] = None
    language: str = "en"


@dataclass
class ExtractedContent:
    content_id: str
    document_id: str
    text_content: str
    tables: List[Dict[str, Any]]
    metadata: Dict[str, Any]
    extraction_mode: ExtractionMode
    extracted_at: str


class DocumentProcessor:
    def __init__(self):
        self.base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        self.documents_dir = os.path.join(self.base_dir, "data", "documents")
        self.metadata_file = os.path.join(self.documents_dir, "metadata.json")
        self.content_file = os.path.join(self.documents_dir, "content.json")
        
        os.makedirs(self.documents_dir, exist_ok=True)
        
        # Load data
        self.documents_metadata = self._load_metadata()
        self.extracted_content = self._load_content()

    def _load_metadata(self) -> Dict[str, DocumentMetadata]:
        """Load document metadata from disk."""
        if os.path.exists(self.metadata_file):
            try:
                with open(self.metadata_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                return {doc_id: DocumentMetadata(**doc) for doc_id, doc in data.items()}
            except Exception:
                pass
        return {}

    def _save_metadata(self):
        """Save document metadata to disk."""
        try:
            data = {doc_id: asdict(doc) for doc_id, doc in self.documents_metadata.items()}
            with open(self.metadata_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2)
        except Exception as e:
            print(f"[DocumentProcessor] Failed to save metadata: {e}")

    def _load_content(self) -> Dict[str, ExtractedContent]:
        """Load extracted content from disk."""
        if os.path.exists(self.content_file):
            try:
                with open(self.content_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                return {content_id: ExtractedContent(**content) for content_id, content in data.items()}
            except Exception:
                pass
        return {}

    def _save_content(self):
        """Save extracted content to disk."""
        try:
            data = {content_id: asdict(content) for content_id, content in self.extracted_content.items()}
            with open(self.content_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2)
        except Exception as e:
            print(f"[DocumentProcessor] Failed to save content: {e}")

    def process_document(self, file_path: str, extraction_mode: ExtractionMode = ExtractionMode.FULL) -> Tuple[DocumentMetadata, ExtractedContent]:
        """
        Process a document and extract content.
        
        Args:
            file_path: Path to the document
            extraction_mode: Mode of extraction
            
        Returns:
            (DocumentMetadata, ExtractedContent)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type
        file_ext = os.path.splitext(file_path)[1].lower()
        doc_type = self._determine_document_type(file_ext)
        
        # Generate document ID
        file_size = os.path.getsize(file_path)
        filename = os.path.basename(file_path)
        document_id = f"doc_{hashlib.md5(f'{filename}_{file_size}'.encode()).hexdigest()[:16]}"
        
        # Extract metadata
        metadata = self._extract_metadata(file_path, doc_type, filename, file_size)
        metadata.document_id = document_id
        
        # Extract content based on type
        if doc_type == DocumentType.PDF:
            content = self._process_pdf(file_path, document_id, extraction_mode)
        elif doc_type == DocumentType.WORD:
            content = self._process_word(file_path, document_id, extraction_mode)
        elif doc_type == DocumentType.EXCEL:
            content = self._process_excel(file_path, document_id, extraction_mode)
        elif doc_type == DocumentType.TEXT:
            content = self._process_text(file_path, document_id, extraction_mode)
        else:
            content = ExtractedContent(
                content_id=f"content_{document_id}",
                document_id=document_id,
                text_content="",
                tables=[],
                metadata={},
                extraction_mode=extraction_mode,
                extracted_at=datetime.now().isoformat()
            )
        
        # Store
        self.documents_metadata[document_id] = metadata
        self.extracted_content[content.content_id] = content
        
        self._save_metadata()
        self._save_content()
        
        return metadata, content

    def _determine_document_type(self, file_ext: str) -> DocumentType:
        """Determine document type from file extension."""
        mapping = {
            '.pdf': DocumentType.PDF,
            '.doc': DocumentType.WORD,
            '.docx': DocumentType.WORD,
            '.xls': DocumentType.EXCEL,
            '.xlsx': DocumentType.EXCEL,
            '.csv': DocumentType.EXCEL,
            '.txt': DocumentType.TEXT,
            '.md': DocumentType.TEXT,
            '.rtf': DocumentType.TEXT,
            '.png': DocumentType.IMAGE,
            '.jpg': DocumentType.IMAGE,
            '.jpeg': DocumentType.IMAGE
        }
        return mapping.get(file_ext, DocumentType.TEXT)

    def _extract_metadata(self, file_path: str, doc_type: DocumentType, 
                        filename: str, file_size: int) -> DocumentMetadata:
        """Extract metadata from document."""
        document_id = f"temp_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        metadata = DocumentMetadata(
            document_id=document_id,
            filename=filename,
            file_type=doc_type,
            file_size=file_size,
            created_at=datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
            processed_at=datetime.now().isoformat(),
            keywords=[]
        )
        
        # Try to extract additional metadata based on type
        try:
            if doc_type == DocumentType.PDF:
                metadata = self._extract_pdf_metadata(file_path, metadata)
            elif doc_type == DocumentType.WORD:
                metadata = self._extract_word_metadata(file_path, metadata)
        except Exception as e:
            print(f"[DocumentProcessor] Metadata extraction failed: {e}")
        
        return metadata

    def _extract_pdf_metadata(self, file_path: str, metadata: DocumentMetadata) -> DocumentMetadata:
        """Extract metadata from PDF file."""
        try:
            import pypdf
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                metadata.page_count = len(reader.pages)
                
                if reader.metadata:
                    if '/Author' in reader.metadata:
                        metadata.author = reader.metadata['/Author']
                    if '/Title' in reader.metadata:
                        metadata.title = reader.metadata['/Title']
                    if '/Subject' in reader.metadata:
                        metadata.subject = reader.metadata['/Subject']
                    if '/Keywords' in reader.metadata:
                        metadata.keywords = reader.metadata['/Keywords'].split(',')
        except ImportError:
            pass
        except Exception as e:
            print(f"[DocumentProcessor] PDF metadata extraction failed: {e}")
        
        return metadata

    def _extract_word_metadata(self, file_path: str, metadata: DocumentMetadata) -> DocumentMetadata:
        """Extract metadata from Word file."""
        try:
            from docx import Document
            doc = Document(file_path)
            
            if doc.core_properties.author:
                metadata.author = doc.core_properties.author
            if doc.core_properties.title:
                metadata.title = doc.core_properties.title
            if doc.core_properties.subject:
                metadata.subject = doc.core_properties.subject
            if doc.core_properties.keywords:
                metadata.keywords = doc.core_properties.keywords.split(',')
            
            # Count pages (approximate)
            metadata.page_count = len(doc.paragraphs) // 20  # Rough estimate
        except ImportError:
            pass
        except Exception as e:
            print(f"[DocumentProcessor] Word metadata extraction failed: {e}")
        
        return metadata

    def _process_pdf(self, file_path: str, document_id: str, 
                    extraction_mode: ExtractionMode) -> ExtractedContent:
        """Process PDF document."""
        text_content = ""
        tables = []
        
        try:
            import pypdf
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                
                for page in reader.pages:
                    text_content += page.extract_text() + "\n"
                
                # Try to extract tables if available
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            page_tables = page.extract_tables()
                            if page_tables:
                                for table in page_tables:
                                    tables.append({
                                        'page': page.page_number,
                                        'data': table
                                    })
                except ImportError:
                    pass
        except ImportError:
            # Fallback: try PyPDF2
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text_content += page.extract_text() + "\n"
            except ImportError:
                text_content = "PDF processing library not available"
        except Exception as e:
            text_content = f"Error processing PDF: {str(e)}"
        
        content_id = f"content_{document_id}"
        
        return ExtractedContent(
            content_id=content_id,
            document_id=document_id,
            text_content=text_content,
            tables=tables,
            metadata={'extraction_method': 'pypdf'},
            extraction_mode=extraction_mode,
            extracted_at=datetime.now().isoformat()
        )

    def _process_word(self, file_path: str, document_id: str,
                     extraction_mode: ExtractionMode) -> ExtractedContent:
        """Process Word document."""
        text_content = ""
        tables = []
        
        try:
            from docx import Document
            doc = Document(file_path)
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({'data': table_data})
        except ImportError:
            text_content = "Word processing library not available"
        except Exception as e:
            text_content = f"Error processing Word document: {str(e)}"
        
        content_id = f"content_{document_id}"
        
        return ExtractedContent(
            content_id=content_id,
            document_id=document_id,
            text_content=text_content,
            tables=tables,
            metadata={'extraction_method': 'python-docx'},
            extraction_mode=extraction_mode,
            extracted_at=datetime.now().isoformat()
        )

    def _process_excel(self, file_path: str, document_id: str,
                      extraction_mode: ExtractionMode) -> ExtractedContent:
        """Process Excel document."""
        text_content = ""
        tables = []
        
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert to text
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += df.to_string()
                text_content += sheet_text + "\n\n"
                
                # Store as table
                tables.append({
                    'sheet': sheet_name,
                    'data': df.values.tolist(),
                    'columns': df.columns.tolist()
                })
        except ImportError:
            text_content = "Excel processing library not available"
        except Exception as e:
            text_content = f"Error processing Excel document: {str(e)}"
        
        content_id = f"content_{document_id}"
        
        return ExtractedContent(
            content_id=content_id,
            document_id=document_id,
            text_content=text_content,
            tables=tables,
            metadata={'extraction_method': 'pandas'},
            extraction_mode=extraction_mode,
            extracted_at=datetime.now().isoformat()
        )

    def _process_text(self, file_path: str, document_id: str,
                     extraction_mode: ExtractionMode) -> ExtractedContent:
        """Process text document."""
        text_content = ""
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
        except UnicodeDecodeError:
            # Try different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text_content = f.read()
            except Exception as e:
                text_content = f"Error reading text file: {str(e)}"
        except Exception as e:
            text_content = f"Error processing text file: {str(e)}"
        
        content_id = f"content_{document_id}"
        
        return ExtractedContent(
            content_id=content_id,
            document_id=document_id,
            text_content=text_content,
            tables=[],
            metadata={'extraction_method': 'native'},
            extraction_mode=extraction_mode,
            extracted_at=datetime.now().isoformat()
        )

    def search_documents(self, query: str, limit: int = 10) -> List[Dict[str, Any]]:
        """
        Search through processed documents.
        
        Args:
            query: Search query
            limit: Maximum results
            
        Returns:
            List of matching documents with snippets
        """
        query_lower = query.lower()
        results = []
        
        for content_id, content in self.extracted_content.items():
            if query_lower in content.text_content.lower():
                # Find snippet
                text = content.text_content
                index = text.lower().find(query_lower)
                start = max(0, index - 50)
                end = min(len(text), index + len(query) + 50)
                snippet = text[start:end]
                
                metadata = self.documents_metadata.get(content.document_id)
                
                results.append({
                    'document_id': content.document_id,
                    'content_id': content_id,
                    'filename': metadata.filename if metadata else 'Unknown',
                    'snippet': snippet,
                    'relevance': text.lower().count(query_lower)
                })
        
        # Sort by relevance
        results.sort(key=lambda x: x['relevance'], reverse=True)
        
        return results[:limit]

    def get_document_content(self, document_id: str) -> Optional[ExtractedContent]:
        """Get extracted content for a document."""
        for content in self.extracted_content.values():
            if content.document_id == document_id:
                return content
        return None

    def get_document_metadata(self, document_id: str) -> Optional[DocumentMetadata]:
        """Get metadata for a document."""
        return self.documents_metadata.get(document_id)

    def delete_document(self, document_id: str) -> bool:
        """Delete a document and its content."""
        if document_id not in self.documents_metadata:
            return False
        
        # Remove metadata
        del self.documents_metadata[document_id]
        
        # Remove content
        to_remove = [content_id for content_id, content in self.extracted_content.items()
                    if content.document_id == document_id]
        for content_id in to_remove:
            del self.extracted_content[content_id]
        
        self._save_metadata()
        self._save_content()
        
        return True

    def get_statistics(self) -> Dict[str, Any]:
        """Get processing statistics."""
        total_documents = len(self.documents_metadata)
        
        # Count by type
        by_type = {}
        for metadata in self.documents_metadata.values():
            doc_type = metadata.file_type.value
            by_type[doc_type] = by_type.get(doc_type, 0) + 1
        
        # Total content size
        total_text_size = sum(len(content.text_content) for content in self.extracted_content.values())
        
        return {
            'total_documents': total_documents,
            'by_type': by_type,
            'total_text_size': total_text_size,
            'total_tables': sum(len(content.tables) for content in self.extracted_content.values())
        }

    def export_content(self, document_id: str, export_path: str) -> Tuple[bool, str]:
        """Export extracted content to file."""
        content = self.get_document_content(document_id)
        if not content:
            return False, "Document not found"
        
        try:
            with open(export_path, 'w', encoding='utf-8') as f:
                f.write(content.text_content)
            
            return True, f"Content exported to {export_path}"
        except Exception as e:
            return False, f"Export failed: {str(e)}"

    def batch_process(self, directory: str, extraction_mode: ExtractionMode = ExtractionMode.FULL) -> Dict[str, Any]:
        """
        Process all documents in a directory.
        
        Args:
            directory: Directory to process
            extraction_mode: Extraction mode
            
        Returns:
            Processing statistics
        """
        if not os.path.exists(directory):
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        processed = 0
        failed = 0
        errors = []
        
        for filename in os.listdir(directory):
            file_path = os.path.join(directory, filename)
            if os.path.isfile(file_path):
                try:
                    self.process_document(file_path, extraction_mode)
                    processed += 1
                except Exception as e:
                    failed += 1
                    errors.append(f"{filename}: {str(e)}")
        
        return {
            'processed': processed,
            'failed': failed,
            'errors': errors
        }
